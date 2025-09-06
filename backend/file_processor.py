import os
import tempfile
import shutil
from typing import List, Optional
from fastapi import UploadFile, HTTPException
import PyPDF2
import docx
import io

class FileProcessor:
    @staticmethod
    async def extract_text_from_pdf(file: UploadFile) -> str:
        """Извлекает текст из PDF файла"""
        try:
            content = await file.read()
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise HTTPException(
                status_code=400, 
                detail=f"Ошибка чтения PDF файла: {str(e)}"
            )

    @staticmethod
    async def extract_text_from_docx(file: UploadFile) -> str:
        """Извлекает текст из DOCX файла"""
        try:
            content = await file.read()
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            raise HTTPException(
                status_code=400, 
                detail=f"Ошибка чтения DOCX файла: {str(e)}"
            )

    @staticmethod
    async def extract_text_from_doc(file: UploadFile) -> str:
        """Извлекает текст из DOC файла (заглушка)"""
        # Для DOC файлов требуется дополнительная обработка
        # В демо-версии возвращаем заглушку
        return "Текст из DOC файла (требуется конвертация)"

    @staticmethod
    async def extract_text_from_file(file: UploadFile) -> str:
        """Определяет тип файла и извлекает текст"""
        filename = file.filename.lower()
        
        if filename.endswith('.pdf'):
            return await FileProcessor.extract_text_from_pdf(file)
        elif filename.endswith('.docx'):
            return await FileProcessor.extract_text_from_docx(file)
        elif filename.endswith('.doc'):
            return await FileProcessor.extract_text_from_doc(file)
        elif filename.endswith('.txt'):
            content = await file.read()
            return content.decode('utf-8', errors='ignore')
        else:
            raise HTTPException(
                status_code=400, 
                detail="Неподдерживаемый формат файла. Поддерживаются: PDF, DOCX, DOC, TXT"
            )

    @staticmethod
    async def process_uploaded_files(files: List[UploadFile]) -> List[str]:
        """Обрабатывает список загруженных файлов и возвращает тексты"""
        texts = []
        for file in files:
            try:
                text = await FileProcessor.extract_text_from_file(file)
                texts.append(text)
            except Exception as e:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Ошибка обработки файла {file.filename}: {str(e)}"
                )
        return texts

    @staticmethod
    async def save_uploaded_file(file: UploadFile, temp_dir: str) -> str:
        """Сохраняет загруженный файл во временную директорию"""
        try:
            file_extension = os.path.splitext(file.filename)[1]
            temp_file_path = os.path.join(temp_dir, f"upload_{os.urandom(8).hex()}{file_extension}")
            
            with open(temp_file_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            
            # Reset file pointer for possible future reads
            await file.seek(0)
            
            return temp_file_path
        except Exception as e:
            raise HTTPException(
                status_code=500, 
                detail=f"Ошибка сохранения файла: {str(e)}"
            )

    @staticmethod
    def read_docx(file_path: str) -> str:
        """Чтение DOCX файла из пути"""
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            raise Exception(f"Ошибка чтения DOCX файла: {str(e)}")

    @staticmethod
    def read_pdf(file_path: str) -> str:
        """Чтение PDF файла из пути"""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"Ошибка чтения PDF файла: {str(e)}")

# Создаем экземпляр для импорта
file_processor = FileProcessor()